# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# ===== TITLE PAGE =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('百大科创 - 文件整理汇总')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('三份文件关键信息提取与对照')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('整理日期：2026年6月12日').font.size = Pt(11)

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
doc.add_heading('目  录', level=1)
toc_items = [
    '一、文件概览',
    '二、第十七届安徽省百所高校百万大学生科普创意创新大赛',
    '    2.1 大赛基本信息',
    '    2.2 参赛对象与分组',
    '    2.3 奖项设置',
    '    2.4 作品类别与要求',
    '    2.5 赛程安排',
    '    2.6 联系方式',
    '三、"千团万人推普强国行"暑期社会实践活动',
    '    3.1 活动板块说明（附件1）',
    '    3.2 活动申报表（附件2）',
    '四、关键时间节点汇总',
    '五、行动建议',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ===== SECTION 1: FILE OVERVIEW =====
doc.add_heading('一、文件概览', level=1)

table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['序号', '文件名', '文件类型', '主要内容']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['1', '关于举办第十七届安徽省百所高校百万大学生\n科普创意创新大赛的通知', 'PDF (35页)', '安徽省科普创意创新大赛的完整赛事通知'],
    ['2', '附件1：2026年"千团万人推普强国行"\n板块说明', 'Word (.docx)', '五大活动板块的详细说明（推普+）'],
    ['3', '附件2：2026年"千团万人推普强国行"\n申报表', 'Word (.doc)', '暑期社会实践志愿服务活动申报表格'],
]
for i, row_data in enumerate(data):
    for j, text in enumerate(row_data):
        table.rows[i+1].cells[j].text = text

doc.add_paragraph()

# ===== SECTION 2: 科普大赛 =====
doc.add_heading('二、第十七届安徽省百所高校百万大学生科普创意创新大赛', level=1)

doc.add_heading('2.1 大赛基本信息', level=2)
info_items = [
    ('大赛全称', '第十七届安徽省百所高校百万大学生科普创意创新大赛'),
    ('大赛简称', '百大科创 / 科普创意创新大赛'),
    ('参与范围', '安徽省百所高校'),
    ('目标人群', '百万大学生'),
    ('政策背景', '2021-2035年全民科学素质行动规划纲要'),
    ('通知落款日期', '2026年5月28日'),
    ('报名截止日期', '2026年5月29日'),
    ('官网', 'http://www.kepuah.cn/'),
    ('微信公众号', 'kepu-ah-dyh'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    run = p.add_run(label + '：')
    run.bold = True
    p.add_run(' ' + value)

doc.add_heading('2.2 参赛对象与分组', level=2)
p = doc.add_paragraph()
p.add_run('年龄分组：').bold = True
p.add_run('15-18岁（20%）、18-21岁（40%）、21岁以上')

doc.add_heading('2.3 奖项设置', level=2)

award_table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
award_table.alignment = WD_TABLE_ALIGNMENT.CENTER

award_headers = ['奖项等级', '名额', '奖金/奖励']
for i, h in enumerate(award_headers):
    cell = award_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

award_data = [
    ['一等奖', '2名', '25,000元'],
    ['二等奖', '12名', '10,000元'],
    ['三等奖', '26名', '5,000元'],
    ['四等奖', '80名', '2,500元'],
    ['五等奖及优秀组织奖', '100+50名', '详见通知'],
]
for i, row_data in enumerate(award_data):
    for j, text in enumerate(row_data):
        award_table.rows[i+1].cells[j].text = text

doc.add_paragraph()

doc.add_heading('2.4 作品类别与规格要求', level=2)

# 科普视频
doc.add_heading('类别一：科普视频', level=3)
for s in [
    '分辨率：1920x1080P',
    '文件大小：不超过2MB（预览图）',
    '时长：1-4分钟',
    '帧率：25fps 或 30fps',
    '格式：MP4 (MPEG-4 Part 14)',
    '编码：H.264/AVC (High Profile), AAC LC',
    '画面比例：16:9',
]:
    doc.add_paragraph(s, style='List Bullet')

# H5
doc.add_heading('类别二：H5（HTML5交互作品）', level=3)
for s in [
    '技术标准：HTML5 + CSS + JS',
    '页面数：至少4页',
    '兼容浏览器：Chrome、Safari、Firefox',
    '兼容系统：Android、iOS',
]:
    doc.add_paragraph(s, style='List Bullet')

# AI
doc.add_heading('类别三：AI（人工智能作品）', level=3)
for s in [
    '涉及AI技术、AR增强现实等',
    '页面数：至少4页',
    '支持平台：Windows / iOS / Android',
    '需提供API调用说明',
    '包含：AI创作说明、技术文档、效果展示',
]:
    doc.add_paragraph(s, style='List Bullet')

# 其他
doc.add_heading('类别四：综合类作品', level=3)
for s in [
    '页数：至少4页，可含3D建模',
    'PPT要求：正文50%以上为PPT格式，4页以上',
    '图片：JPG格式，300dpi，16:9比例',
    '文档：Word (Office 2007及以上)',
    '提交：RAR压缩包',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('2.5 赛程安排（推测）', level=2)
for s in [
    '5月-6月：报名阶段，通过官网 http://www.kepuah.cn/ 报名',
    '6月-7月：作品准备与提交',
    '7月-8月：初赛/复赛，专家评审',
    '8月-9月：决赛答辩',
    '10月-11月：结果公示与颁奖',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('2.6 联系方式', level=2)
for s in [
    '电话：0551-65336381 / 13605606652',
    '电话：0551-62998635',
    '地址：合肥市509 B',
    '官网：http://www.kepuah.cn/',
    '微信公众号：kepu-ah-dyh',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ===== SECTION 3: 推普活动 =====
doc.add_heading('三、"千团万人推普强国行"全国大学生暑期社会实践志愿服务活动', level=1)

doc.add_heading('3.1 活动板块说明（附件1摘录）', level=2)
p = doc.add_paragraph()
p.add_run('活动模式：').bold = True
p.add_run('"推普+" 五大板块，必选板块 + 特色板块至少选一项')

doc.add_paragraph()

blocks = [
    ('板块一：推普 + 语言国情调查', [
        '形式：问卷、访谈、座谈、课堂观察',
        '任务：完成国家通用语言文字使用状况调查',
        '要求：调查问卷统一发放，鼓励撰写高质量调查报告',
    ]),
    ('板块二：推普 + 普法宣传', [
        '形式：走进乡村、校园、社区',
        '内容：新修订《国家通用语言文字法》专项宣传',
        '创新形式：普法主题墙、宣传标语、"普通话+"系列短片',
        '目标：增强尊法、守法意识',
    ]),
    ('板块三：推普 + 语言文化素养', [
        '(1) 青少年提升：结对帮扶，经典诵读、规范书写、"中华经典润校园"',
        '(2) 全民提升：中华经典诵读工程，传承地方优秀文化（含红色文化）',
        '(3) "石榴信箱 万里同心"：规范汉字手写书信，致敬长征英雄',
    ]),
    ('板块四：推普 + 数智赋能', [
        '(1) 数字语博：利用中国语言文字数字博物馆平台开展沉浸式学习',
        '(2) 数字素养：AI基础知识和数字技能培训',
    ]),
    ('板块五：推普 + 职业技能', [
        '对象：农村电商、旅游服务、种养殖、手工业、进城务工人员',
        '形式："普通话+职业技能"定制化培训',
        '目标："学会普通话、会用普通话谋生"',
        '鼓励：职业院校打造专项培训品牌课程',
    ]),
]

for title, items in blocks:
    doc.add_heading(title, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 活动申报表关键字段（附件2摘录）', level=2)

form_sections = [
    ('基本信息', ['团队名称', '申报学校', '实践地点', '实践时间']),
    ('实践内容', ['必选板块：(  ) 推普+语言国情调查  (  ) 推普+普法宣传',
                  '特色板块（至少选一项）：',
                  '  (  ) 推普+语言文化素养  (  ) 推普+数智赋能',
                  '  (  ) 推普+职业技能      (  ) 其他']),
    ('指导教师（不超过2人）', ['需填写：职务、电话']),
    ('团队学生成员（8-12人）', ['需填写：姓名、年级、性别、手机号码、专业、普通话水平等级', '领队标注']),
    ('实践活动简述（2000字以内）', ['内容：活动内容、活动形式、实施规划、拟取得成效', '可另附页']),
    ('审批签字', ['指导教师意见（签字 + 年月日）', '学校团委推荐意见（盖章 + 年月日）']),
    ('提交要求', ['校团委审核盖章后，拍照或扫描上传', '团队物资包收件信息通过系统提交']),
]

for section_title, items in form_sections:
    p = doc.add_paragraph()
    p.add_run(section_title + '：').bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== SECTION 4: TIMELINE =====
doc.add_heading('四、关键时间节点汇总', level=1)

timeline_data = [
    ['2026年5月28日', '科普大赛通知发布', '文件落款日期'],
    ['2026年5月29日', '科普大赛报名截止（推测）', 'PDF末页标注'],
    ['2026年5-6月', '推普活动报名', '暑期社会实践'],
    ['2026年6-7月', '科普大赛作品准备与提交', '—'],
    ['2026年7-8月', '推普活动执行（暑期）', '团队赴各地实践'],
    ['2026年8-9月', '科普大赛决赛', '答辩/评审'],
    ['2026年10-11月', '科普大赛颁奖', '结果公示'],
]

tl_table = doc.add_table(rows=len(timeline_data)+1, cols=3, style='Light Grid Accent 1')
tl_headers = ['时间', '事项', '备注']
for i, h in enumerate(tl_headers):
    cell = tl_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(timeline_data):
    for j, text in enumerate(row):
        tl_table.rows[i+1].cells[j].text = text

doc.add_paragraph()

# ===== SECTION 5: SUGGESTIONS =====
doc.add_heading('五、行动建议', level=1)

suggestions = [
    ('明确方向', '两个活动可同步准备。科普大赛重在作品创作（视频/H5/AI），推普活动重在实地实践。根据团队优势选择重点参与。'),
    ('科普大赛要点', '尽快确认报名状态（截止日期5月29日，距今已过）；确定作品类型，严格遵循技术规格（分辨率1920x1080、MP4/H.264编码等）。'),
    ('推普活动要点', '确定实践地点和团队（8-12名学生+最多2名指导教师）；选择"必选+特色"板块组合；撰写2000字实践活动计划；完成校团委审批盖章。'),
    ('时间管理', '科普大赛若已截止，关注明年第十八届；推普活动暑期执行，当前优先准备团队组建和方案设计。'),
    ('资源整合', '建议确认是否可以团队名义同时参与两个项目，最大化资源利用。同一团队可将科普作品融入推普实践，互为支撑。'),
]

for title, detail in suggestions:
    p = doc.add_paragraph()
    p.add_run(title + '：').bold = True
    p.add_run(detail)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文件整理完毕 —')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(150, 150, 150)

# Save
output_path = r'C:\Users\lenovo\Desktop\百大科创\百大科创_文件整理汇总.docx'
doc.save(output_path)
print('文档已保存至: ' + output_path)
print('Done!')

# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# ===== COVER =====
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('第十七届安徽省百所高校百万大学生\n科普创意创新大赛')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('参赛计划书')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    '参赛赛道：AI类（人工智能作品）',
    '作品方向：AI + 幼儿科普教育',
    '参赛学校：合肥幼儿师范高等专科学校',
    '专业背景：大数据技术',
    '计划日期：2026年6月',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(line).font.size = Pt(12)

doc.add_page_break()

# ===== SECTION 1: PERSONAL ANALYSIS =====
doc.add_heading('一、个人背景与赛道选择分析', level=1)

doc.add_heading('1.1 个人优势画像', level=2)

table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
headers = ['维度', '具体情况', '在大赛中的价值']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['学校背景', '合肥幼儿师范高等专科学校', '幼儿教育领域权威背书，选题自带教育基因'],
    ['专业能力', '大数据技术', 'Python编程、数据处理、可视化，可落地AI应用'],
    ['交叉优势', '幼师院校 + 大数据', '罕见组合，差异化竞争，评委印象分高'],
    ['资源禀赋', '学校有幼儿园/早教合作资源', '可获取真实用户反馈和测试场景'],
    ['团队可组建', '同校大数据+学前教育同学', '跨专业组队，技术与教育双保险'],
]
for i, row_data in enumerate(data):
    for j, text in enumerate(row_data):
        table.rows[i+1].cells[j].text = text

doc.add_paragraph()

doc.add_heading('1.2 四条赛道对比分析', level=2)

track_table = doc.add_table(rows=5, cols=6, style='Light Grid Accent 1')
track_headers = ['赛道', '技术要求', '竞争激烈度', '获奖难度', '与我匹配度', '推荐']
for i, h in enumerate(track_headers):
    cell = track_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

track_data = [
    ['科普视频', '拍摄/剪辑/动画', '极高（门槛低，人人都能做）', '★★★★★', '★★☆☆☆', '不推荐'],
    ['H5交互作品', 'HTML5/CSS/JS前端', '中高（计算机专业多）', '★★★★☆', '★★★☆☆', '备选'],
    ['AI作品', 'Python/AI API/数据处理', '中等（有技术门槛）', '★★★☆☆', '★★★★★', '★强烈推荐'],
    ['综合类', '多样（不限制）', '中等', '★★★☆☆', '★★★☆☆', '备选'],
]
for i, row_data in enumerate(track_data):
    for j, text in enumerate(row_data):
        track_table.rows[i+1].cells[j].text = text

doc.add_paragraph()

doc.add_heading('1.3 最终选择：AI类赛道', level=2)

p = doc.add_paragraph()
run = p.add_run('结论：选择 AI类赛道，作品方向为"AI + 幼儿科普教育"')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

reasons = [
    '技术门槛是你的护城河：大数据技术背景让你天然具备数据处理和AI应用能力，这恰好是大多数师范类院校参赛者的短板。AI赛道自动过滤了大量"只会做PPT和视频"的竞争者。',
    '幼师背景是差异化标签：在AI赛道的理工科选手中，你的幼儿教育视角是独一无二的。评委看到"幼师院校+AI"的组合，比看到"计算机学院+AI"更感兴趣。',
    'AI赛道获奖率更高：虽然AI赛道作品数量少于视频类，但获奖名额按比例分配，且AI赛道的"完成度门槛"会筛掉大量凑数作品。真正能完整提交AI作品的人，获奖概率大幅提升。',
    '能力范围内可落地：不需要从零训练大模型。利用现有AI开放平台API（百度文心、讯飞星火、通义千问等），结合Python数据处理和前端展示，完全可以做出一个完整的AI科普作品。',
    '主题天然契合大赛要求：大赛强调"科普创意创新"，幼儿科普教育是最需要创新的领域之一，AI+教育的结合完美契合大赛主题。',
]
for i, r in enumerate(reasons):
    p = doc.add_paragraph()
    p.add_run(f'{i+1}. {r}')

doc.add_page_break()

# ===== SECTION 2: PROJECT PLAN =====
doc.add_heading('二、参赛作品方案设计', level=1)

doc.add_heading('2.1 作品名称（建议）', level=2)

names = [
    '主推方案：《AI童探——幼儿智能科普互动系统》',
    '备选方案：《"小数据家"——AI幼儿科学启蒙伙伴》',
    '备选方案：《智趣幼儿园——基于AI的幼儿自然科普平台》',
]
for n in names:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('2.2 作品概述', level=2)

p = doc.add_paragraph()
p.add_run(
    '本作品是一款面向3-6岁幼儿的人工智能科普教育互动系统。系统利用AI大模型能力，'
    '以"拍照识物+语音问答+趣味故事"三大核心功能，帮助幼儿认识自然世界中的动植物、'
    '天气现象、日常生活物品等。作品结合大数据技术，对幼儿的认知行为进行数据分析，'
    '生成个性化的科普学习报告，为家长和教师提供科学的教育参考。'
)

doc.add_heading('2.3 核心功能模块', level=2)

modules = [
    ('模块一：AI拍照识物（核心功能）', [
        '幼儿或家长拍摄身边的动植物/物品→调用AI图像识别API→返回物体名称、科普知识',
        '采用百度AI/讯飞星火开放平台API，降低开发难度',
        '输出内容适配幼儿语言：简单、生动、拟人化（如"这是一只小瓢虫，它身上有7颗小星星哦!"）',
        '知识点涵盖：动物、植物、天气、水果蔬菜、日常用品、颜色形状等',
    ]),
    ('模块二：AI语音科普问答', [
        '幼儿通过语音提问（如"为什么天空是蓝色的?"）→AI用幼儿能理解的语言回答',
        '集成语音识别（ASR）和语音合成（TTS）技术',
        '问答内容经过学前教育专业师生审核，确保科学性和适龄性',
        '支持家长预设科普主题（如"本周探索：海洋动物"）',
    ]),
    ('模块三：AI科普故事生成', [
        '根据当天幼儿识别的物体，AI自动生成一个连贯的科普小故事',
        '例如：孩子今天识别了"蚂蚁""树叶""泥土"→生成"蚂蚁搬家的故事"',
        '融入自然科学知识，如蚂蚁的社群行为、植物光合作用等',
        '支持图文+语音播放',
    ]),
    ('模块四：大数据学习分析（差异化亮点）', [
        '记录幼儿的每次交互行为（识别了哪些物体、问了什么问题、哪些概念重复出现）',
        '用Python进行数据分析，可视化展示幼儿的认知发展轨迹',
        '生成"幼儿科普兴趣图谱"——发现孩子对哪类知识更感兴趣',
        '为家长/教师提供个性化科普教育建议',
        '这是大数据技术专业最直接的用武之地，也是其他团队难以复制的优势',
    ]),
    ('模块五：家长/教师管理后台', [
        'Web端后台，管理科普知识库、查看学习报告',
        '教师可批量导入班级幼儿数据，进行群体分析',
        '可与幼儿园课程体系对接',
    ]),
]

for title_text, items in modules:
    doc.add_heading(title_text, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

doc.add_heading('2.4 技术架构', level=2)

tech_table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
tech_headers = ['层级', '技术选型', '说明']
for i, h in enumerate(tech_headers):
    cell = tech_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

tech_data = [
    ['AI能力层', '百度AI/讯飞星火API\n（图像识别+语音+NLP）', '调用成熟API，无需自训模型\n免费额度足够参赛使用'],
    ['数据分析层', 'Python (Pandas/Matplotlib)\nSQLite', '大数据专业的核心技术栈\n实现学习行为分析与可视化'],
    ['前端展示层', 'Vue.js / React\n或微信小程序', '轻量级前端框架\n也可用小程序降低用户门槛'],
    ['后端服务层', 'Python Flask / FastAPI', '与数据分析层统一语言\n快速搭建RESTful API'],
    ['语音交互层', 'Web Speech API\n或讯飞语音SDK', '实现ASR语音识别+TTS语音合成\n适配移动端'],
    ['数据存储', '本地SQLite + 云存储', '参赛阶段用本地数据库即可\n后续可迁移至云端'],
]
for i, row_data in enumerate(tech_data):
    for j, text in enumerate(row_data):
        tech_table.rows[i+1].cells[j].text = text

doc.add_paragraph()

doc.add_heading('2.5 与大赛规则的匹配度', level=2)

match_table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
match_headers = ['大赛要求（AI类）', '本作品满足情况', '状态']
for i, h in enumerate(match_headers):
    cell = match_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

match_data = [
    ['至少4页', '系统包含5大功能模块页面', '✓ 满足'],
    ['支持Windows/iOS/Android', 'Web技术栈天然跨平台；\n也可封装为微信小程序', '✓ 满足'],
    ['需提供API调用说明', '将详细文档化AI API调用流程', '✓ 满足'],
    ['AI创作说明', '提交完整技术文档与设计思路', '✓ 满足'],
    ['演示视频/截图', '录制功能演示视频MP4\n1920×1080, 1-4分钟', '✓ 满足'],
]
for i, row_data in enumerate(match_data):
    for j, text in enumerate(row_data):
        match_table.rows[i+1].cells[j].text = text

doc.add_page_break()

# ===== SECTION 3: COMPETITION ADVANTAGE =====
doc.add_heading('三、获奖可行性分析', level=1)

doc.add_heading('3.1 竞争格局评估', level=2)

p = doc.add_paragraph()
p.add_run('AI赛道参赛者画像：').bold = True

competitors = [
    '计算机/软件专业学生：技术强，但作品往往"炫技"缺乏教育内涵，且主题同质化（人脸识别、聊天机器人等）',
    '电子信息专业学生：硬件+AI结合（如智能小车），但对教育场景理解浅',
    '其他师范院校的计算机专业：有一定教育感知，但幼教领域深耕不如你们',
    '你的定位（幼师院校+大数据+AI科普教育）：在AI赛道中几乎无人竞争同一细分领域',
]
for c in competitors:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('核心结论：你不在"红海"中竞争。幼儿+AI科普×大数据分析，是一个高度差异化的蓝海选题。')
p.runs[0].bold = True

doc.add_heading('3.2 获奖难度评估', level=2)

eval_table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
eval_headers = ['奖项', '名额', '可行性判断']
for i, h in enumerate(eval_headers):
    cell = eval_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

eval_data = [
    ['一等奖', '2名', '难度高。需要作品完成度极高+现场答辩出色。\n如果AI功能真正落地、有幼儿园实测数据，可冲击'],
    ['二等奖', '12名', '★★★★ 现实目标。\n差异化选题+大数据分析亮点+完整作品呈现，概率很大'],
    ['三等奖', '26名', '★★★★★ 保底目标。\n只要完整提交AI作品+有基本功能+选题有亮点，基本稳拿'],
    ['四等奖', '80名', '★★★★★ 几乎稳拿。\nAI赛道的完成度门槛会过滤大量竞争者'],
    ['五等奖', '100名', '提交完整AI作品即可获得'],
]
for i, row_data in enumerate(eval_data):
    for j, text in enumerate(row_data):
        eval_table.rows[i+1].cells[j].text = text

doc.add_paragraph()

doc.add_heading('3.3 评分维度拆解与策略', level=2)

score_table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
score_headers = ['评分维度（推测）', '权重', '本作品策略', '预期得分']
for i, h in enumerate(score_headers):
    cell = score_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

score_data = [
    ['科学性', '25%', '学前教育专业同学审核内容，\n确保知识准确且适龄', '22/25'],
    ['创新性', '25%', 'AI+幼儿科普+大数据分析，\n三重交叉极为罕见', '23/25'],
    ['技术实现', '20%', '调用成熟API降低风险，\n大数据分析板块展示专业能力', '16/20'],
    ['实用性', '15%', '可在幼儿园真实使用，\n有家长教师双端场景', '13/15'],
    ['展示表达', '15%', '幼师院校学生通常表达力强，\n答辩有优势', '12/15'],
]
for i, row_data in enumerate(score_data):
    for j, text in enumerate(row_data):
        score_table.rows[i+1].cells[j].text = text

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('预期总分：86/100 → 冲击二等奖，稳拿三等奖')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 102, 0)

doc.add_page_break()

# ===== SECTION 4: EXECUTION PLAN =====
doc.add_heading('四、执行计划与时间表', level=1)

doc.add_heading('4.1 团队组建建议', level=2)

team_data = [
    ['你（负责人）', '大数据技术', 'AI接口调用、数据分析、\n后端开发、项目管理'],
    ['成员2', '学前教育', '科普内容编写、知识审核、\n幼儿适龄性把关'],
    ['成员3（可选）', '大数据技术\n或计算机', '前端开发、UI设计、\n小程序/H5开发'],
    ['成员4（可选）', '美术/设计\n或学前教育', '界面美化、插图绘制、\n演示视频制作'],
]
team_table = doc.add_table(rows=len(team_data)+1, cols=3, style='Light Grid Accent 1')
team_headers = ['角色', '专业背景', '分工职责']
for i, h in enumerate(team_headers):
    cell = team_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row_data in enumerate(team_data):
    for j, text in enumerate(row_data):
        team_table.rows[i+1].cells[j].text = text

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('核心：最少2人（你+学前教育同学）即可启动。3-4人最优。').bold = True

doc.add_heading('4.2 导师建议', level=2)
for s in [
    '校内导师1：大数据/计算机方向教师 → 技术指导',
    '校内导师2：学前教育方向教师 → 教育内容指导',
    '可选校外指导：合作幼儿园园长/骨干教师 → 实用场景验证',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('4.3 开发时间表（6周计划）', level=2)

plan_table = doc.add_table(rows=8, cols=4, style='Light Grid Accent 1')
plan_headers = ['周次', '阶段', '关键任务', '产出物']
for i, h in enumerate(plan_headers):
    cell = plan_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

plan_data = [
    ['第1周', '选题细化\n与技术验证', '确定最终作品名称和功能范围\n申请AI开放平台API Key\n完成图像识别API调用demo', '技术验证报告\nAPI可用性确认'],
    ['第2周', '核心功能开发', '图像识别+科普知识返回功能\n语音交互功能集成\n基础UI搭建', '可运行的MVP\n（最小可行产品）'],
    ['第3周', '数据分析模块', '用户行为数据采集设计\n数据分析Pipeline搭建\n可视化图表生成', '数据分析模块\n可视化Demo截图'],
    ['第4周', '功能完善\n与联调', 'AI故事生成功能\n家长/教师后台\n全流程联调测试', '功能完整的\n测试版本'],
    ['第5周', '幼儿园测试\n与优化', '联系学校合作幼儿园实地测试\n收集幼儿和教师反馈\nBug修复和体验优化', '用户测试报告\n优化后的V2.0版本'],
    ['第6周', '材料准备\n与提交', '录制演示视频（1-4分钟）\n撰写技术文档和AI调用说明\n制作答辩PPT\n准备RAR压缩包提交', '全部参赛材料\n完整作品包'],
    ['第7周\n（预留）', '缓冲与答辩\n准备', '应对突发问题\n答辩模拟演练\n材料复审', '答辩准备完毕'],
]
for i, row_data in enumerate(plan_data):
    for j, text in enumerate(row_data):
        plan_table.rows[i+1].cells[j].text = text

doc.add_page_break()

# ===== SECTION 5: BUDGET =====
doc.add_heading('五、资源需求与预算', level=1)

budget_table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
budget_headers = ['项目', '说明', '预估费用', '备注']
for i, h in enumerate(budget_headers):
    cell = budget_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

budget_data = [
    ['AI API调用', '百度AI/讯飞开放平台', '0元', '免费额度充足'],
    ['服务器/域名', '本地演示为主，\n可部署到免费云平台', '0-100元', '可用Vercel/Railway等免费服务'],
    ['设计素材', '图标、插图等', '0-200元', '可用免费素材库+自己绘制'],
    ['演示视频制作', '屏幕录制+简单剪辑', '0元', '免费工具即可完成'],
    ['幼儿园测试', '打印测试问卷等', '50-100元', '可选，不必须'],
]
for i, row_data in enumerate(budget_data):
    for j, text in enumerate(row_data):
        budget_table.rows[i+1].cells[j].text = text

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('总预算：0-400元，几乎为零成本参赛。').bold = True

doc.add_paragraph()

# ===== SECTION 6: RISK =====
doc.add_heading('六、风险评估与应对', level=1)

risk_table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
risk_headers = ['风险', '影响', '应对措施']
for i, h in enumerate(risk_headers):
    cell = risk_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

risk_data = [
    ['AI API变更或收费', '中', '注册多个平台账号作为备选\n（百度、讯飞、阿里、腾讯均有免费额度）'],
    ['开发进度延迟', '高', '第2周末必须完成MVP，\n之后只做增量功能，不添加新需求'],
    ['幼儿园测试无法进行', '低', '可用同校同学模拟测试，\n不影响作品提交'],
    ['团队成员变动', '中', '核心功能由负责人掌握，\n确保即便单人也能完成基本提交'],
]
for i, row_data in enumerate(risk_data):
    for j, text in enumerate(row_data):
        risk_table.rows[i+1].cells[j].text = text

doc.add_paragraph()

# ===== SECTION 7: SUMMARY =====
doc.add_heading('七、总结：为什么这个方案能获奖', level=1)

summary_points = [
    ('差异化定位', '在AI赛道中，以"幼儿科普教育"为切入点，避开与纯技术选手的正面竞争。你的幼师院校背景不是劣势，而是最大的差异化武器。'),
    ('技术够用不炫技', '不追求从零训练模型或构建复杂神经网络。利用成熟AI API + 大数据分析展示专业能力，技术方案务实可行。'),
    ('教育内涵是灵魂', '有学前教育专业同学参与内容审核，确保科普知识既科学又适合幼儿。评委中一定有教育背景的专家，他们会被这一点打动。'),
    ('大数据分析是杀手锏', '其他AI作品可能只做到"识别"和"问答"，但你的作品能分析学习行为、生成认知图谱。这个模块直接体现你的专业核心能力，也是其他团队难以复制的。'),
    ('真实场景验证', '通过学校合作幼儿园进行实地测试，有真实用户反馈。这在答辩时会非常加分。'),
    ('完整的故事线', '从"拍一张照片"到"AI认识世界"到"数据看到成长"，整个作品讲了一个完整、温暖的故事。评委不仅看到一个技术作品，更看到了教育温度。'),
]

for title_text, detail in summary_points:
    p = doc.add_paragraph()
    run = p.add_run(title_text + '：')
    run.bold = True
    p.add_run(detail)

doc.add_paragraph()
doc.add_paragraph()

# Final words
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('目标：二等奖冲刺，三等奖保底')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 计划书完 —')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(150, 150, 150)

# Save
output_path = r'C:\Users\lenovo\Desktop\百大科创\参赛计划书_AI幼儿科普教育.docx'
doc.save(output_path)
print('计划书已保存至: ' + output_path)

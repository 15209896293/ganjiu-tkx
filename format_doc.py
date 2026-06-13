#!/usr/bin/env python3
"""
生成格式化后的 团队分工.docx
从原始 .doc 文件提取的内容，重新排版为规范的中文文档。
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── 全局样式配置 ──────────────────────────────────────
FONT_BODY = '宋体'
FONT_HEADING = '黑体'
FONT_CODE = 'Consolas'
SIZE_TITLE = Pt(18)
SIZE_H1 = Pt(15)
SIZE_H2 = Pt(13)
SIZE_BODY = Pt(11)
SIZE_TABLE = Pt(10)
SIZE_SMALL = Pt(9)

COLOR_BLUE = RGBColor(0x1F, 0x4E, 0x79)       # 深蓝
COLOR_HEADER_BG = '1F4E79'                      # 表头背景
COLOR_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)  # 表头文字白色
COLOR_ALT_ROW = 'EBF1F8'                        # 交替行背景
COLOR_BORDER = '808080'                         # 边框灰色
COLOR_CODE_BG = 'F5F5F5'                        # 代码块背景
COLOR_RED = RGBColor(0xCC, 0x00, 0x00)

doc = Document()

# ── 页面设置 ──────────────────────────────────────────
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── 预设样式 ──────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = FONT_BODY
style.font.size = SIZE_BODY
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# ── 辅助函数 ──────────────────────────────────────────

def set_cell_font(cell, font_name=FONT_BODY, size=SIZE_TABLE, bold=False, color=None, east_asian=None):
    """设置单元格内所有 run 的字体"""
    if east_asian is None:
        east_asian = font_name
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = size
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asian)

def add_formatted_paragraph(doc, text, font_name=FONT_BODY, size=SIZE_BODY,
                             bold=False, color=None, alignment=None,
                             space_before=0, space_after=6, line_spacing=1.5,
                             first_line_indent=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    if level == 0:
        # 文档主标题
        p = add_formatted_paragraph(doc, text, font_name=FONT_HEADING, size=SIZE_TITLE,
                                     bold=True, color=COLOR_BLUE,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     space_before=12, space_after=18, line_spacing=1.2)
    elif level == 1:
        # 一级标题（一、二、）
        p = add_formatted_paragraph(doc, text, font_name=FONT_HEADING, size=SIZE_H1,
                                     bold=True, color=COLOR_BLUE,
                                     space_before=18, space_after=10, line_spacing=1.3)
    elif level == 2:
        # 二级标题（1. 2.）
        p = add_formatted_paragraph(doc, text, font_name=FONT_HEADING, size=SIZE_H2,
                                     bold=True, color=COLOR_BLUE,
                                     space_before=14, space_after=8, line_spacing=1.3)
    elif level == 3:
        # 三级标题
        p = add_formatted_paragraph(doc, text, font_name=FONT_HEADING, size=Pt(12),
                                     bold=True,
                                     space_before=10, space_after=6, line_spacing=1.3)
    return p

def add_body(doc, text, indent=True):
    """添加正文段落"""
    return add_formatted_paragraph(doc, text, font_name=FONT_BODY, size=SIZE_BODY,
                                    first_line_indent=0.74 if indent else None)

def add_bold_note(doc, text):
    """添加加粗提示"""
    return add_formatted_paragraph(doc, text, font_name=FONT_BODY, size=SIZE_BODY,
                                    bold=True, space_before=8, space_after=4)

def add_separator(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—' * 30)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    """为整个表格设置统一边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'</w:tblBorders>'
    )
    # Remove existing borders if any
    existing = tblPr.findall(qn('w:tblBorders'))
    for e in existing:
        tblPr.remove(e)
    tblPr.append(borders)

def create_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = FONT_HEADING
        run.font.size = SIZE_TABLE
        run.font.bold = True
        run.font.color.rgb = COLOR_HEADER_TEXT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        set_cell_shading(cell, COLOR_HEADER_BG)

    # 数据行
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.name = FONT_BODY
            run.font.size = SIZE_TABLE
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            # 交替行背景
            if r % 2 == 1:
                set_cell_shading(cell, COLOR_ALT_ROW)

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    set_table_borders(table)
    doc.add_paragraph()  # 表后间距
    return table

def add_code_block(doc, text):
    """添加代码块样式"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    # 添加背景色
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{COLOR_CODE_BG}"/>')
    pPr.append(shading)

    for line in text.split('\n'):
        if line != text.split('\n')[0]:
            p.add_run('\n').font.size = SIZE_SMALL
        run = p.add_run(line)
        run.font.name = FONT_CODE
        run.font.size = SIZE_SMALL
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CODE)

def add_warning(doc, text):
    """添加警告/提醒文本"""
    return add_formatted_paragraph(doc, text, font_name=FONT_BODY, size=SIZE_BODY,
                                    bold=True, color=COLOR_RED,
                                    space_before=6, space_after=6)

def add_bullet(doc, text, level=0):
    """添加项目符号列表"""
    prefix = '  ' * level + '• '
    return add_formatted_paragraph(doc, prefix + text, font_name=FONT_BODY, size=SIZE_BODY,
                                    first_line_indent=None, space_before=2, space_after=2)

def add_cell_with_bold_prefix(cell, bold_text, normal_text):
    """在单元格中添加"加粗前缀 + 普通文本"的段落"""
    cell.text = ''
    p = cell.paragraphs[0]
    run_b = p.add_run(bold_text)
    run_b.font.name = FONT_BODY
    run_b.font.size = SIZE_TABLE
    run_b.font.bold = True
    run_b._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    run_n = p.add_run(normal_text)
    run_n.font.name = FONT_BODY
    run_n.font.size = SIZE_TABLE
    run_n._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


# ══════════════════════════════════════════════════════
#  文档正文开始
# ══════════════════════════════════════════════════════

# ── 主标题 ────────────────────────────────────────────
add_heading_styled(doc, '第十七届"双百大赛"6人团队精细分工方案', level=0)

# ── 引言 ──────────────────────────────────────────────
add_body(doc, '本方案基于你的项目《数据饲养手册：从投喂到驯化，你的第一个智能体》（H5数字科普作品），结合大赛官方要求，设计了6个互补角色，并明确协作流程与审核机制，以保障作品质量、避免脱节、确保按时合规提交。')

add_separator(doc)

# ── 一、6人角色总览 ───────────────────────────────────
add_heading_styled(doc, '一、6人角色总览', level=1)

create_table(doc,
    headers=['角色', '人数', '核心职责', '简要说明'],
    rows=[
        ['技术开发', '1人', '代码实现、交互逻辑、部署上线', '负责H5全部前端代码及链接部署'],
        ['内容+视频+视觉', '1人', '文案、问答、视频、视觉设计', '合并内容、视频、视觉三项工作'],
        ['技术与内容对接', '1人', '需求拆解、可行性评估、功能验收、协调偏差', '防止方案与代码脱节的核心桥梁'],
        ['项目统筹（材料专员）', '1人', '整理材料清单、备份、打包、官网提交', '执行层面的材料准备与提交'],
        ['负责人+审核', '1人', '校内对接、组织测试、最终把关、团队协调', '管理与最终审核'],
        ['复查员', '1人', '独立复查所有材料，输出问题报告', '最终质量把关，不参与执行'],
    ],
    col_widths=[3.0, 1.0, 4.0, 5.5]
)

add_separator(doc)

# ── 二、每个角色的详细任务 ────────────────────────────
add_heading_styled(doc, '二、每个角色的详细任务', level=1)

# ── 2.1 技术开发 ──────────────────────────────────────
add_heading_styled(doc, '1. 技术开发（1人）', level=2)

create_table(doc,
    headers=['任务模块', '具体工作', '输出物'],
    rows=[
        ['代码框架', '搭建 HTML/CSS/JS 结构，确保移动端适配（响应式）', '完整项目文件夹（GitHub 仓库）'],
        ['核心交互', '实现时间轴（点击/滑动）、数据喂养模拟器（滑块+动画）、闯关答题（选择判断）', '可交互的 JS 功能代码'],
        ['实验室模块', '实现 AI 问答模拟（预设回复）或 API 接入（可选）', '可交互的演示功能'],
        ['部署上线', 'GitHub Pages + 腾讯云 EdgeOne Pages 双部署', '2 个可访问链接'],
        ['离线打包', '提供包含所有 HTML/CSS/JS 及资源文件的离线版本', '离线打包文件夹'],
        ['AI 工具声明', '填写《AI工具使用情况说明书》技术部分', '说明书中技术段内容'],
        ['问题修复', '根据测试反馈和复查报告修复 bug', '最终版代码'],
    ],
    col_widths=[2.5, 7.0, 4.0]
)

add_bold_note(doc, '所需技能：会使用 Claude Code 或能独立编写/修改 HTML/CSS/JS，熟悉 Git 基础操作。')

add_separator(doc)

# ── 2.2 内容+视频+视觉 ────────────────────────────────
add_heading_styled(doc, '2. 内容+视频+视觉（1人）', level=2)

create_table(doc,
    headers=['任务模块', '具体工作', '输出物'],
    rows=[
        ['科普文案', '首页引言、AI 时间轴各时代描述、数据旅程步骤说明、实验室引导语', '全部页面文字（Word/腾讯文档）'],
        ['问答设计', '设计 5-8 道闯关题（题干、选项、正确答案、解析），难度递进', '题目库表格'],
        ['教育性设计', '设定学习目标（知识/思维/能力），设计徽章名称（如"数据萌新""驯兽师"）', '徽章文案 + 学习目标'],
        ['申报文档主体', '填写"科学原理及内容、创作目的、设计思路、创新点"', '申报文档 Word（内容部分）'],
        ['引用来源', '整理 AI 发展史引用资料、图片/数据来源', '引用清单'],
        ['演示视频', '使用 OBS 录屏、剪映剪辑、添加片头片尾（含作品名称）、中文字幕', 'MP4 视频（≤4分钟，1920×1080，H.264）'],
        ['视觉设计', '确定配色（科技蓝+白）、页面布局、制作信息图（数据旅程、时间轴）、图标素材', '效果图 JPG（300dpi，16:9）、设计规范、图标包'],
    ],
    col_widths=[2.5, 7.0, 4.0]
)

add_bold_note(doc, '所需技能：文笔好，能用 Canva/PS 做图，会用剪映剪辑，有教育或科普经验者优先。')

add_separator(doc)

# ── 2.3 技术与内容对接 ────────────────────────────────
add_heading_styled(doc, '3. 技术与内容对接（1人）—— 核心新增角色', level=2)

create_table(doc,
    headers=['任务模块', '具体工作', '输出物'],
    rows=[
        ['需求转换', '将内容负责人（第2人）的文案和交互描述拆解成"需求卡片"（功能点、触发方式、预期效果、数据）', '需求卡片表格（Excel/腾讯文档）'],
        ['技术可行性评估', '与技术人员沟通，确定每个功能能否实现、预估工时、风险点', '可行性评估表（标注 P0/P1/P2）'],
        ['任务拆解与排期', '制定功能实现顺序，与技术商定每日交付目标', '开发排期表'],
        ['验收功能', '每完成一个功能，立即测试并反馈内容负责人确认', '功能验收状态表（✅/⚠️/❌）'],
        ['沟通桥梁', '当内容修改需求时，评估影响并通知技术；当技术需要简化时，与内容协商降级方案', '需求变更记录、降级方案文档'],
        ['文档维护', '持续更新需求卡片和验收状态', '最新版协作文档'],
    ],
    col_widths=[3.0, 6.5, 4.0]
)

add_bold_note(doc, '所需技能：既懂一点技术（能判断难度），又能理解方案意图，沟通能力强，有责任心。')

add_separator(doc)

# ── 2.4 项目统筹 ──────────────────────────────────────
add_heading_styled(doc, '4. 项目统筹（材料专员）（1人）', level=2)

create_table(doc,
    headers=['任务模块', '具体工作', '输出物'],
    rows=[
        ['整理材料注意事项清单', '按大赛要求列出所有材料的格式、命名、分辨率、编码、压缩包规范', '《提交材料注意事项清单》（Excel）'],
        ['收集材料', '从各角色处收集最终版代码、图片、视频、文档等', '完整材料文件夹'],
        ['命名检查', '核对所有文件命名是否符合规范', '命名检查记录'],
        ['备份', '将全部材料上传至云盘（百度网盘/阿里云盘）', '云盘备份链接'],
        ['打包压缩', '将所有材料打包为 RAR 压缩包，按大赛规范命名（官网作品编号—作品名称—参赛分类—参赛子分类）', '最终 RAR 文件'],
        ['官网提交', '在大赛官网注册账号、上传 RAR 文件、填写信息', '官网提交成功截图'],
    ],
    col_widths=[3.0, 6.5, 4.0]
)

add_bold_note(doc, '所需技能：细心、有条理、熟悉大赛文件要求，能熟练使用电脑。')

add_separator(doc)

# ── 2.5 负责人+审核 ────────────────────────────────────
add_heading_styled(doc, '5. 负责人+审核（1人）', level=2)

create_table(doc,
    headers=['任务模块', '具体工作', '输出物'],
    rows=[
        ['校内对接', '联系辅导员/科研处，确认本校截止日期、提交方式、所需表格；跑盖章流程（原创承诺书、查重报告）', '校内截止日期、盖章材料'],
        ['测试组织', '在提交前 2 天组织全队进行实测（不同手机、网络），收集问题', '《测试问题反馈表》'],
        ['审核材料', '审核所有材料是否合规（尤其检查是否出现学校/姓名/指导教师信息）', '审核记录'],
        ['最终把关', '确认官网提交和校内提交均已完成，截图留存', '提交证明截图'],
        ['团队协调', '主持每日短会（15分钟），管理进度，处理争议', '会议记录、进度表'],
    ],
    col_widths=[2.5, 7.0, 4.0]
)

add_bold_note(doc, '所需技能：领导力、判断力、熟悉大赛规则，能与学校老师有效沟通。')

add_separator(doc)

# ── 2.6 复查员 ────────────────────────────────────────
add_heading_styled(doc, '6. 复查员（1人）—— 独立质量把关', level=2)

create_table(doc,
    headers=['任务模块', '具体工作', '输出物'],
    rows=[
        ['独立复查', '不参与任何执行工作，只做最终检查', '《复查报告》'],
        ['复查内容',
         '代码：是否含学校/姓名信息（注释、变量名、文件名）\n'
         '视频：片头片尾、时长、字幕、学校信息\n'
         '图片：分辨率、比例、命名、背景简洁\n'
         '文档：模板使用、AI说明书完整、引用齐全\n'
         '压缩包：命名、格式\n'
         '链接：两个链接可正常访问',
         '问题清单及责任人'],
        ['输出报告', '列出所有问题，标记"通过/需修改"，只有复查员签字后负责人才能提交', '签字版复查报告'],
    ],
    col_widths=[2.5, 7.5, 3.5]
)

add_bold_note(doc, '复查报告模板（可直接复制使用）：')

create_table(doc,
    headers=['检查项', '结果（✅/❌）', '问题描述', '责任人', '整改状态'],
    rows=[
        ['代码无学校信息', '✅', '—', '—', '—'],
        ['视频时长≤4分钟', '✅', '—', '—', '—'],
        ['视频有片头片尾+作品名', '✅', '—', '—', '—'],
        ['视频中文字幕', '✅', '—', '—', '—'],
        ['效果图分辨率≥300dpi', '❌', '只有72dpi', '内容+视频+视觉', '已整改'],
        ['效果图比例16:9', '✅', '—', '—', '—'],
        ['压缩包命名规范', '✅', '—', '—', '—'],
        ['离线包包含所有文件', '✅', '—', '—', '—'],
        ['两个链接可访问', '✅', '—', '—', '—'],
        ['申报文档使用官网模板', '✅', '—', '—', '—'],
        ['AI工具说明书完整', '✅', '—', '—', '—'],
        ['引用来源齐全', '⚠️', '缺两张图片来源', '内容+视频+视觉', '已补充'],
        ['无学校/姓名信息（全面）', '✅', '—', '—', '—'],
    ],
    col_widths=[3.5, 2.0, 3.0, 2.5, 2.0]
)

add_formatted_paragraph(doc,
    '复查员必须独立、客观，可邀请不参与本项目的同学或指导老师担任。',
    font_name=FONT_BODY, size=SIZE_BODY, first_line_indent=0.74,
    color=RGBColor(0x55, 0x55, 0x55))

add_separator(doc)

# ── 三、协作流程图 ────────────────────────────────────
add_heading_styled(doc, '三、协作流程图', level=1)

flow_chart = (
    '内容+视频+视觉\n'
    '（产出文案、设计、视频初稿）\n'
    '       ↓\n'
    '技术与内容对接\n'
    '（拆解为需求卡片，评估可行性，排定优先级）\n'
    '       ↓\n'
    '技术开发\n'
    '（按卡片顺序实现功能，每日交付可运行版本）\n'
    '       ↓\n'
    '技术与内容对接\n'
    '（验收功能，反馈内容负责人确认；记录偏差并协调）\n'
    '       ↓\n'
    '负责人+审核\n'
    '（组织全员测试，收集问题，指派修复）\n'
    '       ↓\n'
    '技术开发\n'
    '（修复 bug，优化细节）\n'
    '       ↓\n'
    '项目统筹（材料专员）\n'
    '（收集最终材料，打包，官网提交）\n'
    '       ↓\n'
    '复查员\n'
    '（独立复查，输出报告）\n'
    '       ↓\n'
    '负责人+审核\n'
    '（根据报告整改，最终确认双提交）'
)
add_code_block(doc, flow_chart)

add_separator(doc)

# ── 四、关键协作机制 ──────────────────────────────────
add_heading_styled(doc, '四、关键协作机制', level=1)

# 4.1 每日站会
add_heading_styled(doc, '1. 每日站会（15分钟）', level=2)
add_bullet(doc, '时间：每晚 21:00（线上）')
add_bullet(doc, '参会：全队6人')
add_bullet(doc, '议程：')
add_bullet(doc, '每个人说：今天完成什么？明天做什么？有什么阻碍？', level=1)
add_bullet(doc, '负责人记录进度，调整排期', level=1)

# 4.2 需求卡片
add_heading_styled(doc, '2. 需求卡片表格（技术与内容对接⼈维护）', level=2)
add_body(doc, '格式（可复制到腾讯文档）：')

create_table(doc,
    headers=['功能名称', '位置', '触发方式', '预期效果', '提供数据/文案', '验收标准', '优先级', '状态'],
    rows=[
        ['（示例）', '（示例）', '（示例）', '（示例）', '（示例）', '（示例）', '（示例）', '（示例）'],
    ],
    col_widths=[1.8, 1.2, 1.5, 1.8, 2.0, 1.8, 1.2, 2.0]
)

# 4.3 功能验收
add_heading_styled(doc, '3. 功能验收流程', level=2)
add_bullet(doc, '技术每完成一个功能 → 对接人立即测试 → 录制5秒视频发群里 → 内容负责人确认 → 更新状态为"已完成"')
add_bullet(doc, '若发现偏差 → 对接人记录具体差异 → 与技术协商修改时间 → 更新卡片')

# 4.4 测试与问题反馈
add_heading_styled(doc, '4. 测试与问题反馈', level=2)
add_bullet(doc, '提交前2天，负责人组织全员实测')
add_bullet(doc, '发现问题填写《测试问题反馈表》：')

create_table(doc,
    headers=['问题描述', '发现人', '严重程度（P0/P1/P2）', '责任人', '修复状态'],
    rows=[
        ['（示例问题）', '（示例）', '（示例）', '（示例）', '（示例）'],
    ],
    col_widths=[3.5, 2.0, 3.0, 2.5, 2.0]
)

# 4.5 最终复查
add_heading_styled(doc, '5. 最终提交前复查', level=2)
add_bullet(doc, '复查员拿到所有材料后，逐项检查并填写报告')
add_bullet(doc, '只有复查员签字后，负责人才能执行最终提交')

add_separator(doc)

# ── 五、各阶段时间线 ──────────────────────────────────
add_heading_styled(doc, '五、各阶段时间线（倒推示例）', level=1)

add_body(doc, '假设校内截止 6月30日，官网截止 7月5日。')

create_table(doc,
    headers=['日期', '技术开发', '内容+视频+视觉', '技术与内容对接', '项目统筹', '负责人', '复查员'],
    rows=[
        ['6.20-21', '搭框架、建仓库', '输出文案初稿、视觉风格', '拆解第一轮需求卡片', '整理材料清单', '确认校内截止', '—'],
        ['6.22-23', '实现时间轴+滑块', '完成问答题目、信息图', '验收前两个功能', '—', '组织第一次小测', '—'],
        ['6.24-25', '实现闯关+实验室', '录制视频初稿', '验收新功能、记录偏差', '—', '审核文档初稿', '—'],
        ['6.26-27', '移动端适配+部署', '完成视频剪辑、效果图', '验收全部功能', '收集材料、备份', '组织全员实测', '—'],
        ['6.28-29', '修复测试问题', '最终版文案+视频', '协助验收修复', '打包RAR、官网注册', '审核最终材料', '独立复查'],
        ['6.30', '最终代码提交', '最终材料交付', '确认一切就绪', '官网提交', '校内提交', '签字确认'],
    ],
    col_widths=[1.8, 2.5, 2.5, 2.5, 2.5, 2.5, 1.5]
)

add_separator(doc)

# ── 六、关键提醒 ──────────────────────────────────────
add_heading_styled(doc, '六、关键提醒', level=1)

add_heading_styled(doc, '⚠️ 公正性红线', level=2)
add_warning(doc, '任何地方（代码注释、视频字幕、图片、文档）均不得出现学校、姓名、指导教师信息。违规直接取消资格。')
add_bullet(doc, '复查员必须专门检查这一条。')

add_heading_styled(doc, '⚠️ 材料格式硬性要求', level=2)
add_bullet(doc, '图片：JPG，≥300dpi，16:9')
add_bullet(doc, '视频：MP4，H.264，1920×1080，≤4分钟，片头片尾+中文字幕')
add_bullet(doc, '压缩包：RAR，命名"官网作品编号—作品名称—参赛分类—参赛子分类"')
add_bullet(doc, '申报文档：必须使用大赛官网提供的Word模板')

add_heading_styled(doc, '⚠️ AI工具声明', level=2)
add_warning(doc, '必须在申报文档中增设《AI工具使用情况说明书》，注明工具名称、生成内容、人工修改过程、AI占比。未声明或虚假声明将取消资格。')

add_heading_styled(doc, '⚠️ 双提交缺一不可', level=2)
add_warning(doc, '官网提交（大赛网站）和校内提交（按本校要求）必须都完成。项目统筹负责官网，负责人负责校内。')

add_separator(doc)

# ── 七、可交付的模板清单 ──────────────────────────────
add_heading_styled(doc, '七、可交付的模板清单（需要时可单独提供）', level=1)

add_bullet(doc, '《需求卡片表格》（Excel）')
add_bullet(doc, '《功能验收状态表》')
add_bullet(doc, '《测试问题反馈表》')
add_bullet(doc, '《复查报告模板》')
add_bullet(doc, '《提交材料注意事项清单》（含所有格式规范）')

add_body(doc, '以上内容可直接复制到你的团队文档中使用。如有任何角色职责需要微调，请告知。')

# ── 保存 ──────────────────────────────────────────────
output_path = 'C:/Users/lenovo/Desktop/百大科创/团队分工_排版版.docx'
doc.save(output_path)
print(f'Done: {output_path}')
